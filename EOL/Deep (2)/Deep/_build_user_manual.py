"""One-shot: convert USER_MANUAL.md into a properly formatted .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r"D:\EOL\EOL\Deep (2)\Deep\EOL_MES_User_Manual.docx"

doc = Document()

# ── page setup ───────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def add_heading(text, level=1, color=(31, 73, 125)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(*color)
    if level == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(15)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p


def add_para(text, italic=False, bold=False, size=10.5, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def add_mono(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)
    return p


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_table(headers, rows, col_widths=None, header_color="1F4E79"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    # header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(cell, header_color)
    # body
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = t.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    # spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    return t


# ============================================================
# Title page header
# ============================================================
add_heading("EOL MES + Camera CMS", level=0, color=(31, 73, 125))
add_heading("User Manual", level=0, color=(180, 30, 50))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Toyota Boshoku Device India  •  Bawal Plant  •  YNC SS Line")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(90, 90, 90)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Maintainer: Vivek Kumar    |    Last updated: 2026-05-07")
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(120, 120, 120)


# ============================================================
# 1. Purpose
# ============================================================
add_heading("1. Purpose & Methodology", level=1)

add_heading("What was needed", level=2, color=(60, 100, 50))
add_para(
    "The YNC seat-slider line had to track production, downtime, quality and "
    "process metrics in real time, so operators on the floor and managers in "
    "the office both see the same numbers without anyone manually re-entering "
    "data. We also needed a camera system that records every cycle so any "
    "defect can be replayed."
)

add_heading("Approach used", level=2, color=(60, 100, 50))
add_bullet(
    "One PLC (Mitsubishi Q-series, MC4E protocol on TCP) is the ground-truth "
    "source. It sets bits like L108 = OK, L109 = NG and word registers like "
    "D6005 = status code, D6048 = current model number."
)
add_bullet(
    "A small Python service called the Collector polls the PLC every 30 ms, "
    "debounces the bits, and writes finished cycles into PostgreSQL. The "
    "collector is the only program that talks to the PLC — everything else "
    "just reads the DB."
)
add_bullet(
    "A FastAPI backend turns DB rows into clean JSON for the browser."
)
add_bullet(
    "A React + Vite frontend is the actual screen — operator dashboard, "
    "admin panel, and fullscreen TV view on the shop floor."
)
add_bullet(
    "The Camera CMS is a parallel mini-app: it owns the cameras, records "
    "continuous .ts video per camera, and on demand cuts a clip of any "
    "cycle a user clicks on. It pulls Zone/Line/Machine names from the MES "
    "so both systems show the same hierarchy."
)
add_para(
    "The whole stack is on one PC for now (192.168.10.185). Everything is "
    "reachable over LAN, so any PC on the factory network can open the "
    "dashboard in a browser.", italic=True, size=9.5
)


# ============================================================
# 2. Workflow
# ============================================================
add_heading("2. End-to-End Workflow (PLC → Screen)", level=1)

add_para(
    "The data flow from the physical PLC bit going high to the operator "
    "seeing it on screen takes ~1-2 seconds end to end. The chain:"
)
add_mono(
"""  PLC                  Collector                PostgreSQL
  ─────                ──────────               ──────────
  L108 / L109     ─►   reads bit          ─►   ync_dashboard_complete
  D6005 status    ─►   reads word         ─►   ync_dashboard_complete_ct_log
  D6048 model     ─►   reads word         ─►   ync_status_log
  M-bit (sub)     ─►   edge count         ─►   ync_seatslider
                                                  │
                                                  ▼
                                         FastAPI (port 8080)
                                            GET /api/lines/2/realtime
                                            GET /api/lines/2/ct-history
                                                  │
                                                  ▼
                                         React Frontend (port 5656)
                                            Dashboard / Admin / Fullscreen"""
)

add_para("For video, the chain is parallel:", bold=True, size=10)
add_mono(
"""  Hikvision RTSP camera ─► ffmpeg (24×7) ─► F:\\CameraCMS_Videos\\<cam>\\<shift>.ts
                                                       │
                                                       ▼ (on user click)
                                          ffmpeg -ss/-to clip ─► MP4 player"""
)


# ============================================================
# 3. Directory Structure
# ============================================================
doc.add_page_break()
add_heading("3. Directory Structure", level=1)
add_para('Root project: D:\\EOL\\EOL\\Deep (2)\\Deep\\', bold=True)

add_table(
    ["Path", "Purpose"],
    [
        ["Phase2\\",                            "MES backend (FastAPI). The brains."],
        ["Phase2\\main.py",                     "App entrypoint, DB migrations, CMS proxy endpoints."],
        ["Phase2\\auth.py",                     "JWT login + role checks."],
        ["Phase2\\database.py",                 "Single Postgres connection helper."],
        ["Phase2\\collector_engine.py",         "PLC poller class — only file that reads MC4E bits."],
        ["Phase2\\collectors\\collector_ync_l6.py", "Bootstrap script that launches engine for Line 2."],
        ["Phase2\\routers\\*.py",               "One file per page area: lines, breakdowns, users, quality, etc."],
        ["mes-frontend\\",                      "MES React frontend (Vite)."],
        ["mes-frontend\\src\\pages\\Dashboard.jsx",   "Live shop-floor dashboard."],
        ["mes-frontend\\src\\pages\\Fullscreen.jsx",  "TV display (no login required)."],
        ["mes-frontend\\src\\pages\\AdminPanel.jsx",  "All admin pages: users, lines, machines, cameras."],
        ["mes-frontend\\src\\pages\\Historical.jsx",  "Past-shift analysis."],
        ["mes-frontend\\src\\pages\\ProcessGraphs.jsx", "Sub-machine cycle pulse graphs."],
        ["mes-frontend\\src\\context\\AuthContext.jsx", "Per-tab session + JWT."],
        ["mes-frontend\\src\\api\\client.jsx",  "Axios wrapper used by every page."],
        ["Camera CMS\\",                        "Camera CMS (Flask + React, separate stack)."],
        ["Camera CMS\\backend\\api_server.py",  "Flask API on :5000 + RecordingManager auto-start."],
        ["Camera CMS\\backend\\recorder_engine.py", "ffmpeg-based per-camera continuous .ts recorder."],
        ["Camera CMS\\backend\\cycle_events.py",     "Append-only cycles.csv log."],
        ["Camera CMS\\backend\\plc_poller.py",       "Scaffold — reads L108/L109/M-bit at 30 Hz."],
        ["Camera CMS\\backend\\settings_config.py",  "Stores user-chosen video storage path."],
        ["Camera CMS\\backend\\mes_sync.py",         "One-way pull of zones/lines/machines from MES."],
        ["Camera CMS\\backend\\bin\\ffmpeg.exe",     "Bundled ffmpeg used for record + clip extract."],
        ["Camera CMS\\backend\\zones.json",          "CMS Zone/Line/Machine tree (synced from MES)."],
        ["Camera CMS\\backend\\users.json",          "Local CMS users. Default admin / admin123."],
        ["Camera CMS\\backend\\cycles.csv",          "Append-only cycle log."],
        ["Camera CMS\\frontend\\src\\pages\\masters\\CameraMaster.jsx",     "Camera CRUD with Mounted-On column."],
        ["Camera CMS\\frontend\\src\\pages\\config\\CameraConfig.jsx",     "Bind machine ↔ camera ↔ PLC."],
        ["Camera CMS\\frontend\\src\\pages\\config\\SystemSettings.jsx",   "Storage path + Sync from MES button."],
        ["start_everything.bat",                "THE one launcher. Starts MES + collector + CMS + frontends."],
        ["stop_everything.bat",                 "Kills every Python / Node / ffmpeg process."],
    ],
    col_widths=[7.5, 9.5]
)


# ============================================================
# 4. Network
# ============================================================
doc.add_page_break()
add_heading("4. Network Configuration (IP-by-IP)", level=1)
add_para('Server PC: 192.168.10.185 (also reachable as 127.0.0.1).', bold=True)

add_table(
    ["IP", "Port", "Service", "Notes"],
    [
        ["192.168.10.185", "8080", "MES Backend (FastAPI)",         "All MES API endpoints"],
        ["192.168.10.185", "5656", "MES Frontend (Vite)",           "Operator dashboard, admin"],
        ["192.168.10.185", "5000", "Camera CMS API (Flask)",        "Camera + cycle endpoints"],
        ["192.168.10.185", "8050", "Camera CMS MJPEG Streams",      "Live preview tiles"],
        ["192.168.10.185", "5173", "Camera CMS Frontend (Vite)",    "CMS admin portal"],
        ["192.168.10.210", "5432", "PostgreSQL (energydb)",         "All persistent data. user postgres / tbdi@123"],
        ["192.168.10.150", "5002", "Main PLC (Mitsubishi MC4E)",    "YNC-SS line. L108 = OK, L109 = NG"],
        ["192.168.10.190", "5002", "Sub-PLC #1 — Upper Rail Greasing M/c", "M100 cycle pulse"],
        ["192.168.10.191", "5002", "Sub-PLC #2 — Lock Bar Insert M/c",     "X3 cycle pulse"],
        ["192.168.10.192", "5002", "Sub-PLC #3 — Lower Rail Greasing M/c", "M100 cycle pulse"],
        ["192.168.10.115", "554",  "Hikvision RTSP camera (final station)", "User admin / admin123"],
    ],
    col_widths=[3.2, 1.5, 6.5, 5.5]
)
add_para(
    "Storage: cycle videos go to F:\\CameraCMS_Videos\\ on the external 4 TB "
    "HDD. Path is configurable from CMS → Configuration → System Settings.",
    italic=True, size=9.5
)


# ============================================================
# 5. Page → Data
# ============================================================
add_heading("5. Page → Data → Database Mapping", level=1)
add_para(
    "Each MES page calls a few API endpoints, which read specific tables. "
    "The most important mappings:"
)
add_table(
    ["MES Page", "API Endpoint", "DB Table(s)"],
    [
        ["Dashboard (live OEE, plan vs actual)", "GET /api/lines/2/realtime",                    "ync_dashboard_complete, mes_lines"],
        ["Cycle-Time history graph",             "GET /api/lines/2/ct-history",                  "ync_dashboard_complete_ct_log"],
        ["Loss Distribution panel",              "GET /api/lines/2/hourly-loss-breakdown",       "ync_status_log"],
        ["Shift timeline strip",                 "GET /api/lines/2/status-log",                  "ync_status_log"],
        ["Fullscreen TV (no login)",             "(same endpoints, anonymous)",                  "(same tables)"],
        ["Quality Dashboard / Poka-Yoke",        "GET /api/poka-yoke/live/2",                    "mes_poka_yoke_events, mes_py_master"],
        ["Maintenance Dashboard",                "GET /api/breakdowns/pending-production",       "mes_breakdowns"],
        ["Process Graphs (sub-machines)",        "GET /api/submachines/{id}/hourly",             "mes_machine_process_log"],
        ["Production Hourly History",            "GET /api/lines/2/production_history",          "ync_hourly_production"],
        ["Admin Panel — Cameras list",           "GET /api/cms/cameras (proxy)",                 "via CMS  cameras.json"],
        ["AI Assistant chat",                    "POST /api/ai/chat",                            "reads any of the above"],
    ],
    col_widths=[5.5, 5.5, 5.5]
)

add_para("Camera CMS pages and their stores:", bold=True)
add_table(
    ["CMS Page", "Stored In"],
    [
        ["Camera Master",                "cameras.json (Fernet-encrypted creds)"],
        ["PLC Master",                   "plcs.json"],
        ["Zone / Line / Machine Master", "zones.json (synced from MES Postgres)"],
        ["Camera Config (bindings)",     "camera_config_bindings.json"],
        ["Shift Config",                 "shifts.json"],
        ["Cycle Monitor / Reports",      "cycles.csv (one row per cycle event)"],
        ["System Settings",              "settings.json"],
    ],
    col_widths=[6.5, 10.0]
)


# ============================================================
# 6. How to Start
# ============================================================
doc.add_page_break()
add_heading("6. How to Start the System", level=1)
add_para("One file does it all.", bold=True)

steps = [
    "Plug in the F:\\ external HDD.",
    "Double-click  D:\\EOL\\EOL\\Deep (2)\\Deep\\start_everything.bat",
    "Six command-prompt windows open in this order:",
    "    1. MES-API — uvicorn on :8080",
    "    2. MES-Collector — PLC poll loop, prints [STATUS] IDLE → RUNNING etc.",
    "    3. MES-Frontend — Vite on :5656",
    "    4. CMS-API — Flask on :5000 (also brings up RecordingManager)",
    "    5. CMS-Streams — MJPEG on :8050",
    "    6. CMS-Frontend — Vite on :5173",
    "Browser auto-opens to http://127.0.0.1:5656/",
    "Login with: admin / admin123",
]
for s in steps:
    p = doc.add_paragraph(s)
    p.paragraph_format.space_after = Pt(2)

add_para(
    "To stop everything cleanly: double-click stop_everything.bat", bold=True
)
add_para(
    "If a single service hangs, just close its window; the others keep "
    "running. The collector singleton lock auto-releases after 30 s of no "
    "heartbeat, so a crashed collector can be restarted safely without "
    "manual cleanup.", italic=True, size=9.5
)
add_para("LAN-visible URLs:", bold=True)
add_bullet("MES dashboard: http://192.168.10.185:5656")
add_bullet("CMS portal:    http://192.168.10.185:5173")


# ============================================================
# 7. Roles
# ============================================================
add_heading("7. User Roles & Access", level=1)
add_para(
    "There are six roles. Every login JWT carries the role; the backend "
    "enforces require_role(...) per endpoint and the frontend hides / shows "
    "pages based on the same check, so an operator never even sees an admin "
    "button."
)

add_heading("Role colour scheme (top-banner accent on dashboard)", level=2, color=(60, 100, 50))
add_table(
    ["Role", "Colour", "Used on"],
    [
        ["Admin / Plant Head", "Blue",          "Admin Panel chrome, ADMIN badge"],
        ["Production",         "Green",         "Production dashboard, hourly grid header"],
        ["Maintenance",        "Red",           "Breakdown banner, MTTR/MTBF tiles"],
        ["Quality",            "Amber / Yellow","Poka-yoke live, NG counter pills"],
        ["Operator",           "Grey",          "Default tile, light-mode shop floor"],
        ["Department (generic)","Slate",        "Department-portal landing"],
    ],
    col_widths=[4.5, 3.5, 9.0]
)
add_para(
    "The dashboard layout uses these colours as the four corner accents "
    "(TL = Production green, TR = Quality amber, BL = Maintenance red, "
    "BR = Admin blue) so anyone walking past the screen can tell at a "
    "glance which department's number is healthy.",
    italic=True, size=9.5
)

add_heading("Page access matrix", level=2, color=(60, 100, 50))
add_table(
    ["Page", "Adm", "PH", "Prod", "Maint", "Qual", "Op", "Dept"],
    [
        ["Login",                          "✓",  "✓",  "✓",  "✓",  "✓",  "✓",  "✓"],
        ["Dashboard",                      "✓",  "✓",  "✓",  "✓",  "✓",  "✓",  "✓"],
        ["Fullscreen TV (no login)",       "—",  "—",  "—",  "—",  "—",  "—",  "—"],
        ["Cycle-Time History",             "✓",  "✓",  "✓",  "✓",  "✓",  "RO", "✓"],
        ["Loss Distribution",              "✓",  "✓",  "✓",  "✓",  "✓",  "✓",  "✓"],
        ["Hourly Production Grid",         "✓",  "✓",  "✓",  "✓",  "✓",  "RO", "✓"],
        ["Process Graphs",                 "✓",  "✓",  "✓",  "✓",  "✓",  "RO", "✓"],
        ["Quality Dashboard / Poka-Yoke",  "✓",  "✓",  "RO", "RO", "✓",  "RO", "✓"],
        ["Maintenance Dashboard",          "✓",  "✓",  "RO", "✓",  "RO", "RO", "✓"],
        ["Breakdown raise / fill",         "✓",  "✓",  "P",  "M",  "✗",  "✗",  "✗"],
        ["CAPA / Deviations",              "✓",  "✓",  "✗",  "M",  "Q",  "✗",  "own"],
        ["Admin → Users",                  "✓",  "✓",  "✗",  "✗",  "✗",  "✗",  "✗"],
        ["Admin → Lines / Plants",         "✓",  "✓",  "✗",  "✗",  "✗",  "✗",  "✗"],
        ["Admin → Cameras list",           "✓",  "✓",  "✗",  "✗",  "✗",  "✗",  "✗"],
        ["Admin → Slip Threshold",         "✓",  "✓",  "✗",  "✗",  "✗",  "✗",  "✗"],
        ["AI Assistant chat",              "✓",  "✓",  "✓",  "✓",  "✓",  "✗",  "✓"],
        ["Camera CMS portal",              "✓",  "✓",  "✗",  "✗",  "✗",  "✗",  "✗"],
    ],
    col_widths=[6.0, 1.5, 1.3, 1.5, 1.7, 1.5, 1.3, 1.5]
)
add_para(
    "Legend:  ✓ full   |   RO read-only   |   P production-half-fill   |   "
    "M maintenance-half   |   Q quality-only   |   own = department's own only",
    italic=True, size=9
)
add_para(
    "Admin can also issue per-user, per-page overrides via Admin Panel → "
    "Users → Permissions Matrix (none / read / full columns × every page). "
    "Those overrides win over the default role mapping above."
)

add_heading("Default seeded credentials", level=2, color=(60, 100, 50))
add_table(
    ["User",     "Password",  "Role"],
    [
        ["admin",      "admin123",  "admin"],
        ["supervisor", "super123",  "supervisor / plant_head"],
        ["operator",   "oper123",   "operator"],
    ],
    col_widths=[5.0, 5.0, 7.0]
)
add_para(
    "Change these immediately in production via Admin Panel → Users → "
    "Change Password.", italic=True, size=9.5
)


# ── footer ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "End of manual.   For internal questions contact Vivek Kumar "
    "(Manufacturing & Automation Engineer, TBDI Bawal)."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)


doc.save(OUT)
print(f"WROTE: {OUT}  ({os.path.getsize(OUT)} bytes)")
